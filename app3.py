import os
import streamlit as st
from dotenv import load_dotenv
import pdfplumber
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate

# Load environment variables
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

# Extraction functions for each format:
def get_pdf_text(files):
    text = ""
    for file in files:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
                tables = page.extract_tables()
                for table in tables:
                    text += "\nTable Extracted:\n"
                    for row in table:
                        row_text = " | ".join([cell if cell else "" for cell in row])
                        text += row_text + "\n"
    return text

def get_docx_text(files):
    text = ""
    for file in files:
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            text += "\nTable Extracted:\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text += row_text + "\n"
    return text

def get_csv_text(files):
    text = ""
    for file in files:
        df = pd.read_csv(file)
        text += "\nTable Extracted:\n"
        text += " | ".join(df.columns) + "\n"
        for _, row in df.iterrows():
            row_text = " | ".join([str(cell) for cell in row])
            text += row_text + "\n"
    return text

def get_excel_text(files):
    text = ""
    for file in files:
        xls = pd.ExcelFile(file)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text += f"\nSheet: {sheet_name}\nTable Extracted:\n"
            text += " | ".join(df.columns.astype(str)) + "\n"
            for _, row in df.iterrows():
                row_text = " | ".join([str(cell) for cell in row])
                text += row_text + "\n"
    return text

# Text splitting
def get_text_chunks(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

# Vector store
def get_vector_store(chunks):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

# Conversational chain
def get_conversational_chain():
    prompt_template = """
    Answer the question in detail based on the provided context.
    If the answer is not in the context, say "Answer not found in context."

    Context:
    {context}

    Question:
    {question}

    Answer:
    """
    model = ChatGroq(
        model_name="llama3-8b-8192",
        api_key=api_key,
        temperature=0.3
    )
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(llm=model, chain_type="stuff", prompt=prompt)
    return chain

# User input
def user_input(user_question):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question, k=2)
    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
    return response

# 🚀 Streamlit App
st.set_page_config(page_title="Groq Multi-Format RAG Chatbot 🤖", page_icon="🤖")

# Sidebar
with st.sidebar:
    st.title("EXPLORER:")
    uploaded_files = st.file_uploader(
        "Upload PDF, DOCX, CSV, XLSX/XLS, or TXT files",
        accept_multiple_files=True,
        type=["pdf", "docx", "csv", "xlsx", "xls", "txt"]
    )
    if st.button("Submit & Process"):
        if uploaded_files:
            with st.spinner("Processing..."):
                raw_text = ""
                for file in uploaded_files:
                    if file.name.endswith(".pdf"):
                        raw_text += get_pdf_text([file])
                    elif file.name.endswith(".docx"):
                        raw_text += get_docx_text([file])
                    elif file.name.endswith(".csv"):
                        raw_text += get_csv_text([file])
                    elif file.name.endswith((".xlsx", ".xls")):
                        raw_text += get_excel_text([file])
                    elif file.name.endswith(".txt"):
                        raw_text += file.read().decode("utf-8") + "\n"
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Processing complete! Ask your questions.")
        else:
            st.error("Please upload at least one file.")

# Title
st.title("Chat with Your Files using Groq 🤖")

# Chat memory
if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "assistant", "content": "Upload your files and ask me questions!"}]

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])

# Chat input
prompt = st.chat_input()
if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.write(prompt)

    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            try:
                response = user_input(prompt)
                if response:
                    full_response = response["output_text"]
                    st.write(full_response)
                    st.session_state.messages.append({"role": "assistant", "content": full_response})
                else:
                    st.error("No response generated. Please try again.")
            except Exception as e:
                st.error(f"Error: {e}")